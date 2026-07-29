#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate MPMS Implementation & User Guide (.docx)"""

import os
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUTPUT = os.path.join(os.path.dirname(__file__), "MPMS_Implementation_User_Guide_v2.0.docx")
APP_VERSION = "v2.0"
COMPANY = "PT Karya Machindo Industries"
CREATED = date.today().strftime("%d %B %Y")
GITHUB = "https://github.com/Nazila31/MES-Production-Dashboard.git"


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shd)


def add_field(paragraph, instr_text):
    """Insert a Word field (TOC, PAGE, etc.)."""
    run = paragraph.add_run()
    r = run._r
    for tag, attrs in [
        ("w:fldChar", {"w:fldCharType": "begin"}),
        ("w:instrText", {"xml:space": "preserve"}),
        ("w:fldChar", {"w:fldCharType": "separate"}),
        ("w:fldChar", {"w:fldCharType": "end"}),
    ]:
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        if tag == "w:instrText":
            el.text = instr_text
        r.append(el)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Halaman ")
    run.font.size = Pt(9)
    add_field(p, "PAGE")
    run2 = p.add_run(" dari ")
    run2.font.size = Pt(9)
    add_field(p, "NUMPAGES")


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.bold = True
    run.font.size = Pt(16)
    run.text = "Daftar Isi"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    tp = doc.add_paragraph()
    add_field(tp, r'TOC \o "1-3" \h \z \u')
    note = doc.add_paragraph()
    nrun = note.add_run(
        "Catatan: Buka dokumen ini di Microsoft Word, klik kanan pada Daftar Isi, "
        "lalu pilih \"Perbarui Field\" → \"Perbarui seluruh daftar isi\" agar nomor halaman terisi otomatis."
    )
    nrun.font.size = Pt(9)
    nrun.font.italic = True
    nrun.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def screenshot(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F2F2F2")
    cell.width = Inches(5.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ SCREENSHOT PLACEHOLDER ]\n{caption}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()


def flowchart(doc, lines):
    table = doc.add_table(rows=len(lines), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, line in enumerate(lines):
        cell = table.rows[i].cells[0]
        set_cell_shading(cell, "E8F0FE" if i % 2 == 0 else "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.bold = "↓" not in line and line.strip() != ""
        if line.strip() == "↓":
            run.font.color.rgb = RGBColor(0, 102, 204)
    doc.add_paragraph()


def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "003366")
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def troubleshooting_table(doc, items):
    table = doc.add_table(rows=1 + len(items), cols=3)
    table.style = "Table Grid"
    headers = ["Masalah", "Kemungkinan Penyebab", "Solusi"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "003366")
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
    for ri, (prob, cause, fix) in enumerate(items):
        cells = table.rows[ri + 1].cells
        cells[0].text = prob
        cells[1].text = cause
        cells[2].text = fix
        for c in cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()


# ── cover ────────────────────────────────────────────────────────────────────

def build_cover(doc):
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)

    # Logo placeholder
    logo_table = doc.add_table(rows=1, cols=1)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc = logo_table.rows[0].cells[0]
    set_cell_shading(lc, "F2F2F2")
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run("[ LOGO PERUSAHAAN ]\nPT Karya Machindo Industries")
    lr.font.size = Pt(12)
    lr.font.italic = True
    lr.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Manufacturing Project Monitoring System (MPMS)")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = RGBColor(0, 51, 102)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Implementation & User Guide")
    sr.font.size = Pt(18)
    sr.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        f"Versi Aplikasi: {APP_VERSION}",
        f"Tanggal Pembuatan: {CREATED}",
        COMPANY,
        "Dokumentasi Resmi Implementasi dan Operasional",
    ]:
        mr = meta.add_run(line + "\n")
        mr.font.size = Pt(12)

    doc.add_page_break()


# ── chapters ─────────────────────────────────────────────────────────────────

def chapter_1(doc):
    heading(doc, "BAB 1 – Pendahuluan")
    para(doc, "Manufacturing Project Monitoring System (MPMS) adalah sistem informasi manufaktur yang "
              "dikembangkan untuk PT Karya Machindo Industries guna memantau seluruh siklus proyek "
              "manufaktur — mulai dari penawaran harga (Quotation) hingga pengiriman produk ke pelanggan.")

    heading(doc, "1.1 Tujuan Sistem", 2)
    bullet(doc, "Menyediakan platform terpusat untuk memantau status setiap proyek manufaktur secara real-time.")
    bullet(doc, "Mengintegrasikan alur kerja antar departemen: Marketing, PPIC, Production, dan Admin.")
    bullet(doc, "Mencatat dan mengarsipkan seluruh dokumen proyek (Quotation, SO, SPK, BOM, Work Order, Surat Jalan, Bukti Pengiriman).")
    bullet(doc, "Memberikan laporan produksi dan histori proyek untuk kebutuhan manajemen.")
    bullet(doc, "Mengirim notifikasi otomatis kepada pengguna terkait setiap perubahan status penting.")

    heading(doc, "1.2 Ruang Lingkup Sistem", 2)
    bullet(doc, "Manajemen Quotation dan Follow Up pelanggan (Marketing).")
    bullet(doc, "Pembuatan Sales Order dan pengunggahan dokumen SO/SPK (Admin).")
    bullet(doc, "Perencanaan produksi: BOM, pengecekan stok gudang, Work Order (PPIC).")
    bullet(doc, "Pelaksanaan produksi: Fabrication → Machining → Assembly → Quality Control (Production).")
    bullet(doc, "Pembuatan Surat Jalan dan pengiriman produk (Admin & Production).")
    bullet(doc, "Laporan produksi, histori proyek, dan ekspor data (Admin).")
    bullet(doc, "Sistem notifikasi lintas peran.")

    heading(doc, "1.3 Workflow Bisnis Secara Umum", 2)
    para(doc, "Alur bisnis MPMS mengikuti tahapan berikut:")
    flowchart(doc, [
        "Quotation (Penawaran Harga)",
        "↓",
        "Follow Up Pelanggan",
        "↓",
        "Sales Order (Pesanan Penjualan)",
        "↓",
        "SPK (Surat Perintah Kerja)",
        "↓",
        "Work Order (Perintah Kerja Produksi)",
        "↓",
        "Production (Fabrication → Machining → Assembly → QC)",
        "↓",
        "Quality Control",
        "↓",
        "Surat Jalan",
        "↓",
        "Delivery (Pengiriman)",
        "↓",
        "Reports & Project Finish",
    ])


def chapter_2(doc):
    heading(doc, "BAB 2 – Spesifikasi Sistem")

    heading(doc, "2.1 PC Development", 2)
    para(doc, "PC Development digunakan oleh developer untuk mengembangkan, menguji, dan mempersiapkan sistem sebelum di-deploy ke server produksi.")

    heading(doc, "Software yang Diperlukan", 3)
    sw = [
        ("Git", "Version control — clone dan update project dari GitHub"),
        ("Laragon", "Local development server (Apache/Nginx + PHP + MySQL)"),
        ("PHP 8.3+", "Runtime Laravel (disarankan 8.5 via Laragon)"),
        ("Composer", "Dependency manager PHP"),
        ("MySQL / SQLite", "Database (SQLite default untuk development)"),
        ("HeidiSQL / phpMyAdmin", "Manajemen database visual"),
        ("Visual Studio Code", "Code editor"),
    ]
    simple_table(doc, ["Software", "Fungsi"], sw, [5, 11])

    heading(doc, "Spesifikasi Minimum PC Development", 3)
    simple_table(doc, ["Komponen", "Minimum", "Disarankan"], [
        ("Sistem Operasi", "Windows 10 64-bit", "Windows 10/11 64-bit"),
        ("Prosesor", "Intel Core i3 / setara", "Intel Core i5 atau lebih tinggi"),
        ("RAM", "8 GB", "16 GB"),
        ("Storage", "20 GB ruang kosong", "SSD 50 GB+"),
        ("Browser", "Google Chrome / Microsoft Edge (versi terbaru)", "Google Chrome (versi terbaru)"),
        ("Jaringan", "Akses internet (untuk clone GitHub, Composer)", "Koneksi stabil"),
    ], [4, 5.5, 5.5])

    heading(doc, "2.2 PC Server", 2)
    para(doc, "PC Server menjadi host utama aplikasi MPMS yang diakses oleh seluruh PC Client dan perangkat mobile melalui jaringan LAN/WiFi.")

    heading(doc, "Software PC Server", 3)
    bullet(doc, "Laragon (Apache/Nginx + PHP + MySQL) — atau stack setara")
    bullet(doc, "PHP 8.3+ dengan ekstensi: mbstring, openssl, pdo, tokenizer, xml, ctype, json, fileinfo")
    bullet(doc, "MySQL 8.0+ (disarankan untuk produksi)")
    bullet(doc, "Composer")
    bullet(doc, "Git")

    heading(doc, "Arsitektur Jaringan", 3)
    flowchart(doc, [
        "PC Client / Handphone (Browser)",
        "↓",
        "WiFi / LAN (Jaringan Lokal Perusahaan)",
        "↓",
        "Server MPMS (Laravel + Apache/Nginx)",
        "↓",
        "Database MySQL",
    ])

    heading(doc, "Spesifikasi Minimum PC Server", 3)
    simple_table(doc, ["Komponen", "Minimum", "Disarankan"], [
        ("Sistem Operasi", "Windows Server / Windows 10 Pro", "Windows Server 2019+"),
        ("Prosesor", "Intel Core i5", "Intel Core i7 / Xeon"),
        ("RAM", "8 GB", "16 GB atau lebih"),
        ("Storage", "50 GB SSD", "SSD 256 GB+"),
        ("Jaringan", "Ethernet Gigabit + WiFi Access Point", "Dedicated LAN"),
    ], [4, 5.5, 5.5])

    heading(doc, "2.3 PC Client", 2)
    para(doc, "PC Client adalah komputer pengguna yang hanya mengakses sistem melalui browser web. Tidak diperlukan instalasi software development.")

    bullet(doc, "TIDAK perlu menginstall Visual Studio Code")
    bullet(doc, "TIDAK perlu menginstall Composer")
    bullet(doc, "TIDAK perlu menjalankan Laravel atau PHP")
    bullet(doc, "Cukup membuka browser dan akses: http://IP-SERVER/login.html")
    bullet(doc, "Login menggunakan akun yang diberikan oleh Administrator")

    screenshot(doc, "Halaman login MPMS diakses melalui browser PC Client")

    heading(doc, "2.4 Handphone (Mobile)", 2)
    para(doc, "MPMS dapat diakses melalui browser mobile dengan tampilan responsif (responsive layout).")

    heading(doc, "Browser yang Didukung", 3)
    bullet(doc, "Google Chrome (Android & iOS)")
    bullet(doc, "Safari (iOS)")
    bullet(doc, "Microsoft Edge (Android)")
    bullet(doc, "Firefox Mobile")

    heading(doc, "Cara Membuka Aplikasi", 3)
    numbered(doc, "Pastikan handphone terhubung ke jaringan WiFi yang sama dengan server MPMS.")
    numbered(doc, "Buka browser, ketik alamat: http://IP-SERVER/login.html")
    numbered(doc, "Login dengan akun yang diberikan.")

    heading(doc, "Membuat Shortcut ke Home Screen", 3)
    para(doc, "Android (Chrome):", bold=True)
    bullet(doc, "Buka http://IP-SERVER/login.html di Chrome")
    bullet(doc, "Tap menu (⋮) → \"Add to Home screen\" / \"Tambahkan ke Layar Utama\"")
    bullet(doc, "Beri nama \"MPMS\" → Tap \"Add\"")

    para(doc, "iOS (Safari):", bold=True)
    bullet(doc, "Buka http://IP-SERVER/login.html di Safari")
    bullet(doc, "Tap ikon Share → \"Add to Home Screen\"")
    bullet(doc, "Beri nama \"MPMS\" → Tap \"Add\"")

    heading(doc, "Progressive Web App (PWA) — Rencana Pengembangan", 3)
    para(doc, "Jika MPMS dikembangkan menjadi PWA di versi mendatang, pengguna dapat menginstal aplikasi "
              "seperti aplikasi native melalui prompt \"Install App\" yang muncul di browser, atau melalui "
              "menu browser → \"Install MPMS\". Aplikasi akan berjalan fullscreen tanpa address bar browser.")


def chapter_3(doc):
    heading(doc, "BAB 3 – Instalasi pada PC Development")
    para(doc, "Panduan ini menjelaskan langkah-langkah instalasi MPMS di PC Development menggunakan Laragon di Windows.")
    para(doc, "Semua perintah Laravel dijalankan dari folder backend/, sedangkan perintah Git dijalankan dari root project.")

    steps = [
        ("Clone Project dari GitHub",
         "Buka terminal/PowerShell di folder kerja, lalu clone repository:",
         'git clone https://github.com/Nazila31/MES-Production-Dashboard.git\n'
         'cd "MES-Production-Dashboard"\n'
         '# atau jika folder sudah ada, pull update terbaru:\n'
         'git pull origin main',
         "Terminal PowerShell menampilkan proses clone berhasil"),
        ("Install Dependencies PHP (composer install)",
         "Masuk ke folder backend dan install dependencies:",
         'cd backend\ncomposer install',
         "Output composer install selesai tanpa error fatal"),
        ("Salin File Environment (.env)",
         "Salin file konfigurasi environment:",
         'copy .env.example .env',
         "File .env berhasil dibuat di folder backend"),
        ("Generate Application Key",
         "Generate kunci enkripsi Laravel:",
         'php artisan key:generate',
         "Application key set successfully"),
        ("Konfigurasi Database",
         "Edit file .env — untuk development dengan SQLite (default):",
         'DB_CONNECTION=sqlite\n'
         '# DB_DATABASE otomatis menggunakan database/database.sqlite\n\n'
         '# Untuk MySQL via Laragon:\n'
         'DB_CONNECTION=mysql\n'
         'DB_HOST=127.0.0.1\n'
         'DB_PORT=3306\n'
         'DB_DATABASE=mes_db\n'
         'DB_USERNAME=root\n'
         'DB_PASSWORD=',
         "File .env dengan konfigurasi database"),
        ("Jalankan Migration Database",
         "Buat struktur tabel database:",
         'php artisan migrate',
         "Migration berhasil — semua tabel terbuat"),
        ("Jalankan Database Seeder",
         "Isi data awal (user demo, material sample):",
         'php artisan db:seed\n'
         '# atau sekaligus reset + seed:\n'
         'php artisan migrate:fresh --seed',
         "Database seeded successfully"),
        ("Buat Storage Link",
         "Hubungkan folder upload ke public:",
         'php artisan storage:link',
         "The link has been connected / already exists"),
        ("Jalankan Development Server",
         "Start server Laravel:",
         'php artisan serve',
         "Laravel development server started on http://127.0.0.1:8000"),
        ("Akses Aplikasi",
         "Buka browser dan akses:",
         'http://127.0.0.1:8000/login.html',
         "Halaman login MPMS tampil di browser"),
    ]

    for i, (title, desc, cmd, ss_caption) in enumerate(steps, 1):
        heading(doc, f"3.{i} {title}", 2)
        para(doc, desc)
        code_table = doc.add_table(rows=1, cols=1)
        code_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cc = code_table.rows[0].cells[0]
        set_cell_shading(cc, "F5F5F5")
        cp = cc.paragraphs[0]
        cr = cp.add_run(cmd)
        cr.font.name = "Consolas"
        cr.font.size = Pt(9)
        doc.add_paragraph()
        screenshot(doc, ss_caption)

    heading(doc, "3.11 Akun Demo Setelah Instalasi", 2)
    para(doc, "Setelah db:seed, gunakan akun berikut (password semua: password):")
    simple_table(doc, ["Peran", "Email", "Password"], [
        ("Administrator", "admin@mes.local", "password"),
        ("Marketing", "marketing@mes.local", "password"),
        ("PPIC", "ppic@mes.local", "password"),
        ("Production", "production@mes.local", "password"),
    ], [4, 6, 4])


def chapter_4(doc):
    heading(doc, "BAB 4 – Deploy ke PC Server")
    para(doc, "Panduan deploy MPMS ke PC Server produksi agar dapat diakses oleh seluruh pengguna melalui jaringan LAN/WiFi.")

    heading(doc, "4.1 Clone Project ke Server", 2)
    numbered(doc, "Install Laragon di PC Server (https://laragon.org)")
    numbered(doc, "Install Git for Windows")
    numbered(doc, f"Clone repository: git clone {GITHUB}")
    numbered(doc, "Masuk ke folder backend: cd backend")
    numbered(doc, "Jalankan: composer install --no-dev --optimize-autoloader")
    screenshot(doc, "Clone project dan composer install di PC Server")

    heading(doc, "4.2 Konfigurasi .env Production", 2)
    para(doc, "Salin dan edit .env untuk environment produksi:")
    p = doc.add_paragraph()
    code = (
        "APP_NAME=\"MPMS\"\nAPP_ENV=production\nAPP_DEBUG=false\n"
        "APP_URL=http://192.168.1.100\n\n"
        "DB_CONNECTION=mysql\nDB_HOST=127.0.0.1\nDB_PORT=3306\n"
        "DB_DATABASE=mes_production\nDB_USERNAME=mes_user\nDB_PASSWORD=password_kuat\n\n"
        "SESSION_DRIVER=database\nFILESYSTEM_DISK=public\nLOG_LEVEL=error"
    )
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    para(doc, "PENTING: Set APP_URL ke IP atau hostname server yang akan digunakan client.", bold=True)
    screenshot(doc, "File .env production yang sudah dikonfigurasi")

    heading(doc, "4.3 Setup Database MySQL", 2)
    numbered(doc, "Buka HeidiSQL atau phpMyAdmin via Laragon")
    numbered(doc, "Buat database baru, contoh: mes_production")
    numbered(doc, "Buat user MySQL dengan hak akses penuh ke database tersebut")
    numbered(doc, "Pastikan kredensial sesuai dengan .env")
    screenshot(doc, "Database MySQL mes_production di HeidiSQL")

    heading(doc, "4.4 Migration, Seeder, dan Storage Link", 2)
    p = doc.add_paragraph()
    run = p.add_run(
        "php artisan key:generate\n"
        "php artisan migrate --force\n"
        "php artisan db:seed --force\n"
        "php artisan storage:link"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    heading(doc, "4.5 Konfigurasi Laragon / Apache", 2)
    bullet(doc, "Set document root Laragon ke folder: backend/public")
    bullet(doc, "Buat virtual host, contoh: mpms.test → backend/public")
    bullet(doc, "Untuk akses via IP, pastikan Apache listen di 0.0.0.0:80")
    bullet(doc, "Restart Apache/Nginx setelah konfigurasi")
    screenshot(doc, "Konfigurasi document root Laragon menunjuk ke backend/public")

    heading(doc, "4.6 Pengaturan Firewall Windows", 2)
    numbered(doc, "Buka Windows Defender Firewall → Advanced Settings")
    numbered(doc, "Inbound Rules → New Rule → Port → TCP 80 (HTTP)")
    numbered(doc, "Allow the connection → Apply to Domain, Private")
    numbered(doc, "Ulangi untuk port 443 jika menggunakan HTTPS")

    heading(doc, "4.7 Network Discovery & IP Statis", 2)
    bullet(doc, "Atur IP Statis pada PC Server (contoh: 192.168.1.100)")
    bullet(doc, "Aktifkan Network Discovery di Windows agar server terdeteksi di jaringan")
    bullet(doc, "Catat IP Server — akan digunakan oleh semua client")

    heading(doc, "4.8 Pengujian", 2)
    heading(doc, "Pengujian Localhost (di Server)", 3)
    numbered(doc, "Buka browser di PC Server: http://127.0.0.1/login.html")
    numbered(doc, "Login dengan admin@mes.local / password")
    numbered(doc, "Pastikan dashboard tampil normal")

    heading(doc, "Pengujian via IP Server (dari Client)", 3)
    numbered(doc, "Dari PC Client di jaringan yang sama, buka: http://192.168.1.100/login.html")
    numbered(doc, "Login dan verifikasi semua menu berfungsi")
    screenshot(doc, "Akses MPMS dari PC Client menggunakan IP Server")

    heading(doc, "4.9 Troubleshooting Deploy", 2)
    troubleshooting_table(doc, [
        ("502 Bad Gateway / Blank page", "Document root salah atau PHP tidak jalan", "Pastikan document root = backend/public. Restart Laragon."),
        ("500 Internal Server Error", "APP_KEY belum di-generate atau .env salah", "Jalankan php artisan key:generate. Periksa log di storage/logs/laravel.log"),
        ("Database connection refused", "MySQL tidak jalan atau kredensial salah", "Start MySQL di Laragon. Verifikasi DB_* di .env"),
        ("Storage/file 404", "Symlink belum dibuat", "Jalankan php artisan storage:link"),
        ("Client tidak bisa akses IP", "Firewall memblokir atau AP Isolation", "Buka port 80 di firewall. Nonaktifkan AP Isolation di router WiFi."),
    ])


def chapter_5(doc):
    heading(doc, "BAB 5 – Menghubungkan Laptop Client")
    heading(doc, "5.1 Prasyarat", 2)
    bullet(doc, "PC Client dan PC Server berada pada jaringan yang sama (LAN kabel atau WiFi)")
    bullet(doc, "Server MPMS sudah berjalan dan dapat diakses via localhost")
    bullet(doc, "Firewall server sudah mengizinkan koneksi masuk port 80")

    heading(doc, "5.2 Mengetahui IP Server", 2)
    numbered(doc, "Di PC Server, buka Command Prompt / PowerShell")
    numbered(doc, "Ketik: ipconfig")
    numbered(doc, "Catat IPv4 Address (contoh: 192.168.1.100)")
    screenshot(doc, "Output perintah ipconfig menampilkan IPv4 Address server")

    heading(doc, "5.3 Login dari Browser Client", 2)
    numbered(doc, "Buka browser (Chrome/Edge) di PC Client")
    numbered(doc, "Ketik: http://192.168.1.100/login.html (ganti dengan IP server)")
    numbered(doc, "Masukkan email dan password akun")
    numbered(doc, "Klik Login — sistem akan mengarahkan ke dashboard sesuai peran")
    screenshot(doc, "Login berhasil dari PC Client")

    heading(doc, "5.4 Troubleshooting Koneksi Client", 2)
    troubleshooting_table(doc, [
        ("Browser: \"Can't reach this page\"", "Server mati atau IP salah", "Pastikan Laragon/Apache jalan. Verifikasi IP dengan ipconfig."),
        ("Connection timed out", "Firewall memblokir", "Buka port 80 di Windows Firewall server."),
        ("Login berhasil tapi data kosong", "API tidak terjangkau", "Pastikan APP_URL di .env server sesuai IP yang diakses client."),
        ("WiFi client tidak bisa akses", "AP Isolation aktif di router", "Nonaktifkan AP Isolation / Client Isolation di pengaturan router WiFi."),
        ("Halaman tampil tapi CSS/JS error", "Mixed content atau path salah", "Akses via http:// bukan https://. Clear cache browser."),
    ])


def chapter_6(doc):
    heading(doc, "BAB 6 – Menggunakan MPMS melalui Handphone")
    heading(doc, "6.1 Membuka Aplikasi", 2)
    numbered(doc, "Hubungkan handphone ke WiFi perusahaan (sama dengan server)")
    numbered(doc, "Buka browser → http://IP-SERVER/login.html")
    numbered(doc, "Login dengan akun yang diberikan")
    screenshot(doc, "Tampilan login MPMS di handphone")

    heading(doc, "6.2 Responsive Layout", 2)
    para(doc, "MPMS dirancang responsif — sidebar otomatis menjadi menu hamburger di layar kecil. "
              "Tabel dapat di-scroll horizontal. Semua fitur utama tetap dapat diakses.")

    heading(doc, "6.3 Shortcut Home Screen", 2)
    para(doc, "Lihat panduan di BAB 2.4 untuk Android dan iOS.")

    heading(doc, "6.4 Saran Penggunaan Mobile", 2)
    bullet(doc, "Gunakan orientasi portrait untuk navigasi, landscape untuk tabel lebar")
    bullet(doc, "Pastikan koneksi WiFi stabil — MPMS membutuhkan koneksi ke server lokal")
    bullet(doc, "Upload dokumen dari handphone didukung (PDF, JPG, PNG max 10 MB)")
    bullet(doc, "Logout setelah selesai jika menggunakan perangkat bersama")


def chapter_7(doc):
    heading(doc, "BAB 7 – Panduan Login")
    heading(doc, "7.1 Login", 2)
    numbered(doc, "Buka http://IP-SERVER/login.html")
    numbered(doc, "Masukkan Email (contoh: admin@mes.local)")
    numbered(doc, "Masukkan Password")
    numbered(doc, "Klik tombol Login")
    numbered(doc, "Sistem mengarahkan ke dashboard sesuai peran pengguna")
    screenshot(doc, "Form login MPMS")

    heading(doc, "7.2 Logout", 2)
    numbered(doc, "Klik nama/avatar profil di pojok kanan atas header")
    numbered(doc, "Pilih \"Logout\" / \"Keluar\"")
    numbered(doc, "Sistem menghapus token sesi dan kembali ke halaman login")

    heading(doc, "7.3 Lupa Password", 2)
    para(doc, "Fitur reset password mandiri belum tersedia di versi saat ini. "
              "Hubungi Administrator sistem untuk reset password melalui database atau seeder ulang.")

    heading(doc, "7.4 Hak Akses Setiap Role", 2)
    simple_table(doc, ["Peran", "Email Demo", "Akses Menu"], [
        ("Administrator", "admin@mes.local", "Semua menu: Dashboard, Quotations, Sales Orders, PPIC, Production, Delivery, Reports, Notifications"),
        ("Marketing", "marketing@mes.local", "Dashboard Marketing, Quotations, Notifications"),
        ("PPIC", "ppic@mes.local", "Dashboard PPIC, PPIC (BOM/Warehouse/WO/Schedule), Notifications"),
        ("Production", "production@mes.local", "Dashboard Production, Production, Delivery, Notifications"),
    ], [3, 4.5, 8.5])


def chapter_8(doc):
    heading(doc, "BAB 8 – Panduan Setiap Role")

    # Admin
    heading(doc, "8.1 Administrator (Admin)", 2)
    para(doc, "Admin memiliki akses penuh ke seluruh modul sistem.")

    admin_menus = [
        ("Dashboard", "Ringkasan KPI: total quotation, SO aktif, proyek per tahap produksi, proyek terlambat, tren 7 hari, deadline reminders, activity feed."),
        ("Quotations", "Lihat, buat, edit, approve/reject quotation. Kelola follow-up pelanggan. Upload dokumen quotation."),
        ("Sales Orders", "Buat SO dari quotation approved. Upload dokumen SO dan SPK. Atur deadline material & produksi. Buat Surat Jalan setelah QC passed."),
        ("PPIC", "Akses penuh modul perencanaan: BOM, Warehouse, Work Order, Schedule."),
        ("Production", "Monitor dan operasikan tahap produksi. QC pass/reject."),
        ("Delivery", "Pantau dan kelola pengiriman produk."),
        ("Reports", "Laporan produksi lengkap, filter, export CSV/PDF, histori proyek, preview dokumen."),
        ("Notifications", "Semua notifikasi sistem — quotation approved, SO created, WO released, stage completed, QC passed, ready for delivery, project completed."),
    ]
    simple_table(doc, ["Menu", "Fungsi"], admin_menus, [3.5, 12.5])
    screenshot(doc, "Dashboard Administrator dengan KPI dan grafik")

    # Marketing
    heading(doc, "8.2 Marketing", 2)
    para(doc, "Marketing bertanggung jawab atas seluruh proses penawaran harga dan follow-up pelanggan.")

    mkt_features = [
        ("Dashboard Marketing", "Statistik quotation (draft/sent/approved/rejected), quotation terbaru, panel follow-up, tren 7 hari."),
        ("Buat Quotation", "Form: nomor quotation, klien, PIC, mesin, nominal, deskripsi. Upload dokumen (PDF/gambar). Status awal: Draft."),
        ("Edit Quotation", "Edit quotation berstatus Draft atau Sent. Quotation Approved/Rejected terkunci."),
        ("Follow Up", "Catat follow-up: tanggal, deskripsi, status (Menunggu Respon, Negosiasi, Revisi, Disetujui, Ditolak)."),
        ("Approve Quotation", "Ubah status quotation menjadi Approved — siap dibuat Sales Order oleh Admin."),
        ("Reject Quotation", "Tolak quotation dengan status Rejected."),
        ("Notifications", "Terima notifikasi terkait quotation dan aktivitas sistem."),
    ]
    simple_table(doc, ["Fitur", "Deskripsi"], mkt_features, [4, 12])
    screenshot(doc, "Halaman daftar Quotation — role Marketing")

    # PPIC
    heading(doc, "8.3 PPIC (Production Planning & Inventory Control)", 2)
    para(doc, "PPIC mengelola perencanaan produksi setelah Sales Order dibuat Admin.")

    ppic_features = [
        ("Dashboard PPIC", "Antrian planning (waiting PPIC / processing), panel deadline, tren aktivitas planning."),
        ("Released SO List", "Daftar SO menunggu perencanaan PPIC."),
        ("BOM Input", "Input Bill of Materials: kode material, nama, qty, satuan. Upload dokumen produksi. Status SO → ppic_processing."),
        ("Warehouse Stock", "Cek ketersediaan stok material di gudang vs kebutuhan BOM."),
        ("Work Order", "Buat Work Order (draft), upload file WO, release ke lantai produksi. Status SO → in_production."),
        ("Production Schedule", "Lihat daftar Work Order yang sudah released beserta tanggal rilis."),
        ("Atur Deadline", "Set/edit deadline material dan deadline produksi pada Sales Order."),
        ("Notifications", "Notifikasi SO baru dibuat, dll."),
    ]
    simple_table(doc, ["Fitur", "Deskripsi"], ppic_features, [4, 12])
    screenshot(doc, "Halaman BOM Input — role PPIC")

    # Production
    heading(doc, "8.4 Production (Operator Produksi)", 2)
    para(doc, "Production menjalankan tahap produksi dan pengiriman produk.")

    prod_features = [
        ("Dashboard Production", "Proyek in-production per tahap, panel delivery tasks, tren produksi."),
        ("Production Monitor", "Daftar proyek aktif dengan progress bar per tahap (Fabrication 25%, Machining 50%, Assembly 75%, QC 100%)."),
        ("Start Stage", "Mulai tahap produksi — catat waktu mulai dan operator."),
        ("Finish Stage", "Selesaikan tahap — otomatis lanjut ke tahap berikutnya."),
        ("QC Pass", "Selesaikan QC → status SO menjadi qc_passed. Admin dapat buat Surat Jalan."),
        ("QC Reject", "Tolak QC: isi alasan, pilih tahap kembali (Fabrication/Machining/Assembly). Proyek kembali ke tahap tersebut."),
        ("Delivery", "Input data pengiriman: nomor resi, kurir, upload bukti pengiriman. Status → completed."),
        ("Notifications", "Notifikasi WO released, ready for delivery."),
    ]
    simple_table(doc, ["Fitur", "Deskripsi"], prod_features, [4, 12])
    screenshot(doc, "Halaman Production dengan progress tahap dan tombol QC Reject")


def chapter_9(doc):
    heading(doc, "BAB 9 – Workflow Lengkap")
    para(doc, "Bab ini menjelaskan alur kerja detail dari awal hingga proyek selesai, termasuk skenario QC Reject.")

    heading(doc, "9.1 Diagram Alur Lengkap", 2)
    flowchart(doc, [
        "1. Marketing: Buat Quotation",
        "↓",
        "2. Marketing: Follow Up Pelanggan (opsional, berulang)",
        "↓",
        "3. Marketing/Admin: Approve Quotation",
        "↓",
        "4. Admin: Buat Sales Order (+ upload SO & SPK)",
        "↓",
        "5. PPIC: Input BOM + Cek Warehouse",
        "↓",
        "6. PPIC: Buat & Release Work Order",
        "↓",
        "7. Production: Fabrication → Start & Finish",
        "↓",
        "8. Production: Machining → Start & Finish",
        "↓",
        "9. Production: Assembly → Start & Finish",
        "↓",
        "10. Production: Quality Control",
        "↓",
        "11a. QC PASS → Admin: Buat Surat Jalan",
        "↓",
        "12. Production: Delivery (input resi + bukti)",
        "↓",
        "13. Project Finish → Reports",
    ])

    heading(doc, "9.2 Alur QC Reject", 2)
    flowchart(doc, [
        "QC Inspection — TEMUKAN CACAT",
        "↓",
        "Production: Klik \"Reject QC\"",
        "↓",
        "Isi Alasan Reject + Pilih Tahap Kembali",
        "↓",
        "Fabrication / Machining / Assembly",
        "↓",
        "Kerjakan ulang tahap tersebut → Finish",
        "↓",
        "Lanjut tahap berikutnya hingga QC",
        "↓",
        "QC Ulang → Pass → Lanjut Surat Jalan",
    ])

    heading(doc, "9.3 Penjelasan Setiap Tahap", 2)
    stages = [
        ("Quotation", "Marketing", "Buat penawaran harga, upload dokumen, follow-up pelanggan."),
        ("Approve", "Marketing/Admin", "Quotation disetujui — status approved, siap jadi SO."),
        ("Sales Order", "Admin", "Buat SO dari quotation approved. Upload SO & SPK. SO status: waiting_ppic."),
        ("BOM", "PPIC", "Input daftar material. SO status: ppic_processing."),
        ("Work Order", "PPIC", "Buat WO, upload file, release. SO status: in_production."),
        ("Fabrication", "Production", "Tahap fabrikasi (25%). Start → Finish."),
        ("Machining", "Production", "Tahap machining (50%). Start → Finish."),
        ("Assembly", "Production", "Tahap assembly (75%). Start → Finish."),
        ("QC", "Production", "Quality Control (100%). Pass atau Reject."),
        ("Surat Jalan", "Admin", "Buat delivery note. SO status: ready_for_delivery."),
        ("Delivery", "Production", "Input resi & bukti kirim. SO status: completed."),
    ]
    simple_table(doc, ["Tahap", "Pelaku", "Keterangan"], stages, [3, 3, 10])


def chapter_10(doc):
    heading(doc, "BAB 10 – Reports")
    para(doc, "Modul Reports tersedia khusus untuk Administrator. Menyediakan laporan produksi komprehensif.")

    heading(doc, "10.1 Ringkasan KPI", 2)
    bullet(doc, "Total Orders — jumlah seluruh Sales Order")
    bullet(doc, "Completed — proyek selesai")
    bullet(doc, "Delayed — proyek terlambat (deadline produksi terlewati & belum selesai)")
    bullet(doc, "Efficiency — persentase penyelesaian tepat waktu")

    heading(doc, "10.2 Filter Laporan", 2)
    bullet(doc, "Rentang tanggal (Date From — Date To)")
    bullet(doc, "Klien / Client")
    bullet(doc, "Nomor Sales Order")
    bullet(doc, "Nomor Quotation")
    bullet(doc, "Status Proyek")
    screenshot(doc, "Panel filter Reports dengan berbagai opsi")

    heading(doc, "10.3 Grafik", 2)
    bullet(doc, "Monthly Production — tren produksi 6 bulan terakhir (line chart)")
    bullet(doc, "Department Distribution — distribusi tahap: Fabrication, Machining, Assembly, QC (doughnut chart)")
    bullet(doc, "Status Distribution — Completed, In Production, Waiting PPIC (bar chart)")

    heading(doc, "10.4 Project History", 2)
    para(doc, "Tabel histori proyek menampilkan:")
    simple_table(doc, ["Kolom", "Keterangan"], [
        ("Quotation Date", "Tanggal quotation dibuat"),
        ("SO Date", "Tanggal Sales Order dibuat"),
        ("SO Number", "Nomor Sales Order"),
        ("Client", "Nama klien/pelanggan"),
        ("Material Deadline", "Deadline ketersediaan material (dengan indikator warna)"),
        ("Production Deadline", "Deadline produksi (dengan indikator warna)"),
        ("Production Start", "Tanggal mulai produksi"),
        ("Completion Date", "Tanggal selesai"),
        ("Total Hari", "Durasi total proyek (hari)"),
        ("Status", "Status terkini SO"),
        ("Documents", "Tombol expand — lihat semua dokumen proyek"),
    ], [4, 12])

    heading(doc, "10.5 Documents (Expandable)", 2)
    para(doc, "Klik baris proyek untuk expand detail:")
    bullet(doc, "Semua dokumen terunggah: Quotation, SO, SPK, BOM, Work Order, Surat Jalan, Bukti Delivery")
    bullet(doc, "Preview PDF — buka di modal/ tab baru")
    bullet(doc, "Preview Image — tampilkan gambar JPG/PNG")
    bullet(doc, "Download — unduh file ke perangkat")
    bullet(doc, "QC Reject History — riwayat penolakan QC jika ada")

    heading(doc, "10.6 Export", 2)
    bullet(doc, "Export CSV — unduh Production_Report.csv untuk analisis di Excel")
    bullet(doc, "Export PDF — unduh Production_Report.pdf (format laporan formal)")
    screenshot(doc, "Tombol Export CSV dan Export PDF di halaman Reports")


def chapter_11(doc):
    heading(doc, "BAB 11 – Notifikasi")
    para(doc, "MPMS mengirim notifikasi otomatis kepada pengguna terkait setiap peristiwa penting dalam workflow.")

    heading(doc, "11.1 Badge Notifikasi", 2)
    para(doc, "Ikon lonceng (bell) di header kanan menampilkan badge angka merah — jumlah notifikasi belum dibaca.")

    heading(doc, "11.2 Dropdown Notifikasi", 2)
    para(doc, "Klik ikon lonceng untuk membuka dropdown berisi notifikasi terbaru. Klik notifikasi untuk navigasi ke halaman terkait.")

    heading(doc, "11.3 Halaman Notifikasi", 2)
    para(doc, "Menu Notifications menampilkan seluruh notifikasi dengan opsi mark as read / mark all as read.")

    heading(doc, "11.4 Jenis Notifikasi", 2)
    simple_table(doc, ["Tipe", "Pemicu", "Penerima", "Navigasi"], [
        ("quotation_approved", "Quotation disetujui", "Admin", "Halaman Quotations"),
        ("so_created", "Sales Order dibuat", "PPIC", "Halaman PPIC / SO Detail"),
        ("work_order_released", "Work Order dirilis", "Production", "Halaman Production"),
        ("stage_completed", "Tahap produksi selesai", "Admin", "Halaman Production"),
        ("qc_passed", "QC lulus", "Admin", "SO Detail"),
        ("ready_for_delivery", "Surat Jalan dibuat", "Production", "SO Detail"),
        ("project_completed", "Pengiriman selesai", "Admin", "Halaman Reports"),
    ], [3.5, 4, 2.5, 5])

    heading(doc, "11.5 Status Read / Unread", 2)
    bullet(doc, "Unread — notifikasi belum dibaca (badge counter aktif, background highlighted)")
    bullet(doc, "Read — notifikasi sudah dibaca (badge berkurang)")
    bullet(doc, "Mark All Read — tandai semua sekaligus sudah dibaca")
    screenshot(doc, "Dropdown notifikasi dengan badge dan daftar notifikasi")


def chapter_12(doc):
    heading(doc, "BAB 12 – Troubleshooting")
    troubleshooting_table(doc, [
        ("Laravel tidak jalan / artisan not found", "Perintah dijalankan dari folder root, bukan backend/", "cd backend lalu jalankan php artisan serve"),
        ("Database gagal connect", "MySQL tidak jalan atau .env salah", "Start MySQL di Laragon. Periksa DB_HOST, DB_DATABASE, DB_USERNAME, DB_PASSWORD di .env"),
        ("Storage link error / file 404", "Symlink belum ada", "php artisan storage:link. Pastikan document root = backend/public"),
        ("PDF tidak bisa dibuka / preview kosong", "File corrupt atau path salah", "Re-upload dokumen. Periksa APP_URL di .env. Cek storage/app/public/"),
        ("Upload gagal", "File terlalu besar atau format tidak didukung", "Max 10 MB. Format: PDF, JPG, JPEG, PNG saja"),
        ("Login gagal / 401 Unauthorized", "Email/password salah atau token expired", "Periksa kredensial. Clear sessionStorage browser. Coba logout & login ulang"),
        ("Permission denied / 403 Forbidden", "Role tidak memiliki akses ke modul", "Pastikan login dengan role yang benar. Admin akses semua modul"),
        ("404 Not Found pada halaman", "URL salah atau file tidak ada", "Pastikan URL benar: /login.html, /pages/... . Periksa file di backend/public/"),
        ("Notifikasi tidak muncul", "API notifications error atau JS conflict", "Refresh halaman. Periksa Console browser (F12). Pastikan token auth valid"),
        ("QC Reject tidak berjalan", "Modal/form error atau stage salah", "Pastikan berada di tahap QC. Isi alasan reject dan pilih return stage"),
        ("CSS/JS tidak load", "Path asset salah", "Akses via server Laravel (port 8000), bukan buka file HTML langsung"),
        ("Migration error", "Tabel sudah ada atau DB belum dibuat", "php artisan migrate:fresh --seed (HATI-HATI: hapus semua data)"),
        ("CORS / API error di client", "APP_URL tidak sesuai", "Set APP_URL di .env server = URL yang diakses client"),
    ])


def chapter_13(doc):
    heading(doc, "BAB 13 – FAQ (Pertanyaan Umum)")

    faqs = [
        ("Bagaimana cara reset password?",
         "Fitur reset password mandiri belum tersedia. Hubungi Administrator untuk reset via database "
         "(update kolom password di tabel users menggunakan bcrypt hash) atau jalankan ulang seeder."),
        ("Bagaimana cara upload dokumen?",
         "Pada form Quotation, Sales Order, BOM, Work Order, Surat Jalan, atau Delivery — klik tombol "
         "Browse/Upload, pilih file (PDF/JPG/PNG max 10 MB), lalu simpan form."),
        ("Bagaimana jika QC Reject?",
         "Di halaman Production, pada proyek yang sedang di tahap QC, klik \"Reject QC\". Isi alasan "
         "penolakan dan pilih tahap kembali (Fabrication/Machining/Assembly). Proyek akan kembali ke "
         "tahap tersebut untuk dikerjakan ulang."),
        ("Bagaimana melihat Reports?",
         "Login sebagai Admin → menu Reports. Gunakan filter untuk mempersempit data. Klik baris proyek "
         "untuk expand dokumen. Gunakan tombol Export CSV/PDF."),
        ("Bagaimana mengakses dari laptop lain?",
         "Pastikan laptop terhubung ke jaringan yang sama dengan server. Buka browser → http://IP-SERVER/login.html → login."),
        ("Bagaimana mengakses dari handphone?",
         "Hubungkan ke WiFi perusahaan. Buka browser → http://IP-SERVER/login.html → login. "
         "Buat shortcut Home Screen untuk akses cepat."),
        ("Bagaimana update sistem dari GitHub?",
         "Di root project: git pull origin main. Di backend: composer install → php artisan migrate → "
         "php artisan optimize. Restart server."),
        ("Bagaimana backup database?",
         "MySQL: gunakan HeidiSQL → Export database. Atau: mysqldump -u user -p mes_production > backup.sql. "
         "SQLite: salin file backend/database/database.sqlite."),
        ("Apakah bisa akses dari luar kantor?",
         "Saat ini MPMS dirancang untuk jaringan lokal (LAN/WiFi). Akses dari luar membutuhkan VPN "
         "atau port forwarding — konsultasikan dengan IT."),
        ("Format file apa saja yang didukung?",
         "PDF, JPG, JPEG, PNG — maksimal 10 MB per file."),
    ]

    for i, (q, a) in enumerate(faqs, 1):
        heading(doc, f"13.{i} {q}", 2)
        para(doc, a)


def chapter_14(doc):
    heading(doc, "BAB 14 – Maintenance")
    para(doc, "Panduan pemeliharaan rutin sistem MPMS untuk Administrator Server.")

    heading(doc, "14.1 Backup Database", 2)
    para(doc, "MySQL (via HeidiSQL):", bold=True)
    numbered(doc, "Buka HeidiSQL → connect ke mes_production")
    numbered(doc, "Klik kanan database → Export Database as SQL")
    numbered(doc, "Simpan file backup dengan tanggal, contoh: mes_backup_2026-07-29.sql")
    para(doc, "MySQL (via command line):", bold=True)
    p = doc.add_paragraph()
    run = p.add_run("mysqldump -u mes_user -p mes_production > D:\\Backup\\mes_backup.sql")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    para(doc, "Disarankan: backup otomatis mingguan.")

    heading(doc, "14.2 Update Project dari GitHub", 2)
    p = doc.add_paragraph()
    run = p.add_run(
        'cd "D:\\path\\to\\MES_Cleaned_Project - Copy"\n'
        "git pull origin main\n"
        "cd backend\n"
        "composer install --no-dev\n"
        "php artisan migrate --force\n"
        "php artisan optimize"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    heading(doc, "14.3 Composer Install & Update", 2)
    bullet(doc, "composer install — install dependencies sesuai composer.lock")
    bullet(doc, "composer update — update dependencies ke versi terbaru (hati-hati di production)")

    heading(doc, "14.4 Migration", 2)
    bullet(doc, "php artisan migrate — jalankan migration baru")
    bullet(doc, "php artisan migrate:status — cek status migration")
    bullet(doc, "php artisan migrate:rollback — rollback batch terakhir (hati-hati)")

    heading(doc, "14.5 Optimasi & Cache", 2)
    p = doc.add_paragraph()
    run = p.add_run(
        "php artisan optimize       # cache config + routes\n"
        "php artisan config:cache   # cache konfigurasi\n"
        "php artisan route:cache    # cache routes\n"
        "php artisan view:cache     # cache views\n"
        "php artisan cache:clear    # bersihkan cache\n"
        "php artisan config:clear   # bersihkan config cache\n"
        "php artisan route:clear    # bersihkan route cache"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    heading(doc, "14.6 Restart Server", 2)
    numbered(doc, "Stop Apache/Nginx di Laragon (Stop All)")
    numbered(doc, "Start All di Laragon")
    numbered(doc, "Verifikasi akses http://IP-SERVER/login.html")

    heading(doc, "14.7 Monitoring Log Laravel", 2)
    para(doc, "Log aplikasi tersimpan di: backend/storage/logs/laravel.log")
    bullet(doc, "Periksa log saat terjadi error 500 atau masalah tidak terduga")
    bullet(doc, "Rotate log berkala agar tidak memenuhi disk")
    bullet(doc, "Set LOG_LEVEL=error di .env production untuk mengurangi noise")


def appendix(doc):
    heading(doc, "Lampiran")

    heading(doc, "A. Struktur Folder Project", 2)
    p = doc.add_paragraph()
    run = p.add_run(
        "MES_Cleaned_Project/          ← Root Git repository\n"
        "├── README.md\n"
        "├── backend/                    ← Laravel API + Frontend\n"
        "│   ├── app/\n"
        "│   │   ├── Enums/              ← Status enums\n"
        "│   │   ├── Http/Controllers/   ← REST API controllers\n"
        "│   │   ├── Models/             ← Eloquent models\n"
        "│   │   ├── Services/           ← Notification, activity\n"
        "│   │   └── Support/            ← File rules, storage URLs\n"
        "│   ├── database/\n"
        "│   │   ├── migrations/         ← Database migrations\n"
        "│   │   └── seeders/            ← Demo data seeder\n"
        "│   ├── public/                 ← Web root (document root)\n"
        "│   │   ├── index.html          ← Admin dashboard\n"
        "│   │   ├── login.html\n"
        "│   │   ├── pages/              ← Feature pages\n"
        "│   │   └── assets/             ← CSS, JS, images\n"
        "│   ├── routes/                 ← api.php, web.php\n"
        "│   └── storage/app/public/     ← Uploaded files\n"
        "└── docs/                       ← Dokumentasi"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    heading(doc, "B. Struktur Database", 2)
    simple_table(doc, ["Tabel", "Fungsi"], [
        ("users", "Data pengguna + role"),
        ("quotations", "Data quotation/penawaran"),
        ("quotation_follow_ups", "Riwayat follow-up quotation"),
        ("sales_orders", "Sales Order + status + deadlines + file paths"),
        ("materials", "Master data material gudang"),
        ("bom_items", "Bill of Materials per SO"),
        ("work_orders", "Work Order per SO"),
        ("production_stage_logs", "Log setiap tahap produksi"),
        ("qc_reject_logs", "Riwayat penolakan QC"),
        ("activities", "Activity feed/log"),
        ("notifications", "Notifikasi pengguna"),
        ("personal_access_tokens", "Token autentikasi Sanctum"),
    ], [4.5, 11.5])

    heading(doc, "C. Daftar Role dan Hak Akses", 2)
    simple_table(doc, ["Role", "Menu", "API Access"], [
        ("admin", "Semua menu", "Semua endpoint"),
        ("marketing", "Dashboard, Quotations, Notifications", "Quotations CRUD, approve/reject, follow-ups"),
        ("ppic", "Dashboard, PPIC, Notifications", "PPIC BOM/Warehouse/WO, SO deadlines"),
        ("production", "Dashboard, Production, Delivery, Notifications", "Production stages, QC, shipment"),
    ], [3, 5, 8])

    heading(doc, "D. Daftar Status Project", 2)
    heading(doc, "Status Quotation", 3)
    simple_table(doc, ["Status", "Label", "Keterangan"], [
        ("draft", "Draft", "Baru dibuat, dapat diedit"),
        ("sent", "Sent", "Sudah dikirim ke pelanggan"),
        ("approved", "Approved", "Disetujui — siap jadi SO"),
        ("rejected", "Rejected", "Ditolak — terkunci"),
    ], [3, 3, 10])

    heading(doc, "Status Sales Order", 3)
    simple_table(doc, ["Status", "Label", "Trigger"], [
        ("waiting_ppic", "Waiting PPIC", "SO baru dibuat"),
        ("ppic_processing", "PPIC Processing", "BOM disimpan"),
        ("in_production", "In Production", "Work Order released"),
        ("qc_passed", "QC Passed", "QC selesai (pass)"),
        ("ready_for_delivery", "Ready for Delivery", "Surat Jalan dibuat"),
        ("completed", "Completed", "Pengiriman selesai"),
    ], [3.5, 3.5, 9])

    heading(doc, "Indikator Deadline", 3)
    simple_table(doc, ["Status", "Label UI", "Kondisi"], [
        ("safe", "Aman", "Lebih dari 7 hari sebelum deadline"),
        ("approaching", "Mendekati", "7 hari atau kurang sebelum deadline"),
        ("overdue", "Terlewati", "Sudah melewati deadline"),
    ], [3, 3, 10])

    heading(doc, "E. Daftar Jenis Dokumen Upload", 2)
    simple_table(doc, ["Dokumen", "Folder Storage", "Diupload Oleh", "Format"], [
        ("Quotation", "quotations/", "Marketing/Admin", "PDF, JPG, PNG"),
        ("Sales Order", "sales-orders/", "Admin", "PDF, JPG, PNG"),
        ("SPK", "spk/", "Admin", "PDF, JPG, PNG"),
        ("BOM / Production Doc", "bom/", "PPIC", "PDF, JPG, PNG"),
        ("Work Order", "work-orders/", "PPIC", "PDF, JPG, PNG"),
        ("Surat Jalan", "delivery-notes/", "Admin", "PDF, JPG, PNG"),
        ("Bukti Delivery/Resi", "shipment-proofs/", "Production", "PDF, JPG, PNG"),
    ], [3.5, 3, 3, 5.5])
    para(doc, "Semua file maksimal 10 MB per upload.")

    heading(doc, "F. Diagram Relasi Workflow", 2)
    flowchart(doc, [
        "Quotation ──→ Follow Up ──→ Approve",
        "↓",
        "Sales Order (+ SO File + SPK File)",
        "↓",
        "BOM Items ←── Materials (Warehouse)",
        "↓",
        "Work Order ──→ Release",
        "↓",
        "Production Stage Logs",
        "(Fabrication → Machining → Assembly → QC)",
        "↓",
        "QC Reject Logs (jika reject → kembali ke stage)",
        "↓",
        "Surat Jalan → Shipment Proof → Completed",
        "↓",
        "Reports & Activities & Notifications",
    ])


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    build_cover(doc)
    add_toc(doc)

    chapter_1(doc)
    doc.add_page_break()
    chapter_2(doc)
    doc.add_page_break()
    chapter_3(doc)
    doc.add_page_break()
    chapter_4(doc)
    doc.add_page_break()
    chapter_5(doc)
    doc.add_page_break()
    chapter_6(doc)
    doc.add_page_break()
    chapter_7(doc)
    doc.add_page_break()
    chapter_8(doc)
    doc.add_page_break()
    chapter_9(doc)
    doc.add_page_break()
    chapter_10(doc)
    doc.add_page_break()
    chapter_11(doc)
    doc.add_page_break()
    chapter_12(doc)
    doc.add_page_break()
    chapter_13(doc)
    doc.add_page_break()
    chapter_14(doc)
    doc.add_page_break()
    appendix(doc)

    # Page numbers on all sections
    for sec in doc.sections:
        add_page_number_footer(sec)
        # Different first page (cover has no number ideally)
        sec.different_first_page_header_footer = True

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Guidebook saved: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
